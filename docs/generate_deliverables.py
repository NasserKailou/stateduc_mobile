#!/usr/bin/env python3
"""
generate_deliverables.py — Génération des livrables Word (.docx) StatEduc Mobile
Auteur : Abdoul Nasser Kailou
Projet : PAQABU / UNESCO — Ministère de l'Éducation Nationale du Burundi
Version : 1.0 — Juillet 2026

Usage :
    pip install python-docx
    cd docs/
    python generate_deliverables.py

Les fichiers .docx sont générés dans docs/output/
"""

import os
import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERREUR : python-docx n'est pas installé.")
    print("Installez-le avec : pip install python-docx")
    sys.exit(1)

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────

AUTHOR     = "Abdoul Nasser Kailou"
PROJECT    = "PAQABU / UNESCO — Ministère de l'Éducation Nationale du Burundi"
DATE       = "Juillet 2026"
REPO_URL   = "https://github.com/NasserKailou/stateduc_mobile"

# Répertoire racine du projet (parent de docs/)
ROOT_DIR   = Path(__file__).parent.parent
OUTPUT_DIR = Path(__file__).parent / "output"

# Couleurs institutionnelles
COLOR_TITLE    = RGBColor(0x00, 0x38, 0x7A)   # Bleu MEN
COLOR_HEADING1 = RGBColor(0x00, 0x38, 0x7A)
COLOR_HEADING2 = RGBColor(0x1F, 0x5C, 0x9E)
COLOR_HEADING3 = RGBColor(0x2E, 0x74, 0xB5)
COLOR_CODE_BG  = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_TABLE_H  = RGBColor(0x00, 0x38, 0x7A)   # En-tête tableau
COLOR_TABLE_HR = RGBColor(0xFF, 0xFF, 0xFF)   # Texte en-tête tableau

# ─────────────────────────────────────────────────────────────────────────────
# CATALOGUE DES LIVRABLES
# ─────────────────────────────────────────────────────────────────────────────

DELIVERABLES = [
    {
        "id":       "presentation",
        "title":    "StatEduc Mobile — Présentation Technique Complète",
        "subtitle": "Architecture du système et composants techniques",
        "source":   "stateduc_flutter/presentation.md",
        "output":   "01_Presentation_Technique_StatEduc_Mobile.docx",
        "category": "Document de présentation",
    },
    {
        "id":       "architecture",
        "title":    "StatEduc Mobile — Architecture Technique",
        "subtitle": "Guide développeur Flutter et architecture applicative",
        "source":   "stateduc_flutter/architecture_technique.md",
        "output":   "02_Architecture_Technique_Flutter.docx",
        "category": "Document technique",
    },
    {
        "id":       "administration",
        "title":    "StatEduc Mobile — Guide d'Administration",
        "subtitle": "Manuel utilisateur et administrateur A→Z",
        "source":   "administration.md",
        "output":   "03_Guide_Administration.docx",
        "category": "Guide opérationnel",
    },
    {
        "id":       "restitution",
        "title":    "StatEduc Mobile — Restitution Technique",
        "subtitle": "Bilan des travaux de développement",
        "source":   "RESTITUTION_TECHNIQUE_STATEDUC_MOBILE.md",
        "output":   "04_Restitution_Technique.docx",
        "category": "Compte-rendu technique",
    },
    {
        "id":       "notepresentation",
        "title":    "StatEduc Mobile — Note de Présentation",
        "subtitle": "Synthèse pour transfert de compétences",
        "source":   "notepresentation.md",
        "output":   "05_Note_Presentation.docx",
        "category": "Synthèse exécutive",
    },
    {
        "id":       "recapitulatif",
        "title":    "StatEduc Mobile — Récapitulatif Technique",
        "subtitle": "Guide développeur et mainteneur",
        "source":   "recapitulatif.md",
        "output":   "06_Recapitulatif_Technique.docx",
        "category": "Guide développeur",
    },
    {
        "id":       "analysis",
        "title":    "StatEduc — Analyse de l'Application Existante",
        "subtitle": "Analyse technique de l'ancienne application Cordova",
        "source":   "ANALYSIS.md",
        "output":   "07_Analyse_Application.docx",
        "category": "Analyse technique",
    },
    {
        "id":       "changelog_flutter",
        "title":    "StatEduc Mobile — Journal des Modifications Flutter",
        "subtitle": "Historique des versions de l'application Flutter",
        "source":   "stateduc_flutter/CHANGELOG.md",
        "output":   "08_Journal_Modifications_Flutter.docx",
        "category": "Journal des modifications",
    },
    {
        "id":       "changelog_server",
        "title":    "StatEduc — Journal des Modifications Serveur PHP",
        "subtitle": "Historique des versions du serveur PHP",
        "source":   "StatEduc_burundi/CHANGELOG.md",
        "output":   "09_Journal_Modifications_Serveur.docx",
        "category": "Journal des modifications",
    },
    {
        "id":       "release_signing",
        "title":    "StatEduc Mobile — Procédure de Signature APK Release",
        "subtitle": "Guide de build et signature de l'APK Android",
        "source":   "stateduc_flutter/RELEASE_SIGNING.md",
        "output":   "10_Procedure_Signature_Release.docx",
        "category": "Guide opérationnel",
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# UTILITAIRES
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_background(cell, fill_color: RGBColor):
    """Colorer le fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{fill_color[0]:02X}{fill_color[1]:02X}{fill_color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_page_break(doc: Document):
    """Ajouter un saut de page."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def setup_document_styles(doc: Document):
    """Configurer les styles du document."""
    styles = doc.styles

    # Style Normal
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_HEADING1

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_HEADING2

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_HEADING3

    # Heading 4
    try:
        h4 = styles['Heading 4']
        h4.font.name = 'Calibri'
        h4.font.size = Pt(11)
        h4.font.bold = True
        h4.font.italic = True
        h4.font.color.rgb = COLOR_HEADING3
    except Exception:
        pass


def add_title_page(doc: Document, deliverable: dict):
    """Ajouter une page de titre professionnelle."""
    # Espace supérieur
    for _ in range(4):
        doc.add_paragraph()

    # Catégorie
    p_cat = doc.add_paragraph()
    p_cat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cat = p_cat.add_run(deliverable["category"].upper())
    run_cat.font.name = 'Calibri'
    run_cat.font.size = Pt(10)
    run_cat.font.color.rgb = COLOR_HEADING2
    run_cat.font.bold = True

    # Séparateur
    doc.add_paragraph()

    # Titre principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(deliverable["title"])
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_TITLE

    # Sous-titre
    doc.add_paragraph()
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(deliverable["subtitle"])
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = COLOR_HEADING2
    run_sub.font.italic = True

    # Ligne de séparation (simulation via paragraphe avec bordure)
    for _ in range(3):
        doc.add_paragraph()

    # Bloc informations
    info_lines = [
        f"Auteur : {AUTHOR}",
        f"Projet : {PROJECT}",
        f"Date : {DATE}",
        f"Dépôt : {REPO_URL}",
    ]
    for line in info_lines:
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_info = p_info.add_run(line)
        run_info.font.name = 'Calibri'
        run_info.font.size = Pt(10.5)
        run_info.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Saut de page après la page de titre
    add_page_break(doc)


# ─────────────────────────────────────────────────────────────────────────────
# PARSEUR MARKDOWN → DOCX
# ─────────────────────────────────────────────────────────────────────────────

def parse_inline(text: str) -> list:
    """
    Analyser le texte inline Markdown et retourner une liste de (texte, bold, italic, code).
    """
    result = []
    # Pattern: **bold**, *italic*, `code`, ~~strikethrough~~
    pattern = re.compile(
        r'(`[^`]+`)'           # backtick code
        r'|(\*\*[^*]+\*\*)'   # bold **...**
        r'|(\*[^*]+\*)'       # italic *...*
        r'|(_[^_]+_)'         # italic _..._
        r'|(~~[^~]+~~)'       # strikethrough
    )

    last = 0
    for m in pattern.finditer(text):
        # Texte avant le match
        if m.start() > last:
            result.append((text[last:m.start()], False, False, False, False))

        matched = m.group(0)
        if matched.startswith('`') and matched.endswith('`'):
            result.append((matched[1:-1], False, False, True, False))
        elif matched.startswith('**') and matched.endswith('**'):
            result.append((matched[2:-2], True, False, False, False))
        elif matched.startswith('*') and matched.endswith('*'):
            result.append((matched[1:-1], False, True, False, False))
        elif matched.startswith('_') and matched.endswith('_'):
            result.append((matched[1:-1], False, True, False, False))
        elif matched.startswith('~~') and matched.endswith('~~'):
            result.append((matched[2:-2], False, False, False, True))
        last = m.end()

    if last < len(text):
        result.append((text[last:], False, False, False, False))

    return result


def apply_inline_formatting(para, text: str):
    """Appliquer le formatage inline à un paragraphe."""
    tokens = parse_inline(text)
    for (content, bold, italic, code, strike) in tokens:
        if not content:
            continue
        run = para.add_run(content)
        run.font.name = 'Courier New' if code else 'Calibri'
        run.font.size = Pt(9.5) if code else None
        run.bold = bold
        run.italic = italic
        if code:
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)  # rouge foncé pour code inline
        if strike:
            run.font.strike = True


def add_table_from_md(doc: Document, table_lines: list):
    """Créer un tableau Word à partir de lignes Markdown."""
    # Filter separator lines (|---|---|)
    data_lines = [l for l in table_lines if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', l)]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        # Split by | and strip
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ''
            # Remove bold markers from header cells (they're already styled)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            if is_header:
                run.bold = True
                run.font.color.rgb = COLOR_TABLE_HR
                set_cell_background(cell, COLOR_TABLE_H)
            else:
                run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

    doc.add_paragraph()  # espace après le tableau


def md_to_docx(doc: Document, md_content: str):
    """
    Convertir le contenu Markdown en éléments Word.
    Supporte :
      - Titres # ## ### ####
      - Listes - * + et numérotées
      - Blocs de code ```
      - Tableaux |col|col|
      - Paragraphes normaux avec formatage inline
      - Séparateurs ---
      - Blockquotes >
    """
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    # Skip the H1 title (already on the title page)
    skip_first_h1 = True

    while i < len(lines):
        line = lines[i]

        # ── Code block ──────────────────────────────────────────────────────
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
                    p.paragraph_format.left_indent = Cm(0.8)
                    # Background-like style via shading
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    pPr.append(shd)
                    code_lines = []
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table detection ──────────────────────────────────────────────────
        if line.strip().startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            # Check if next line is still table
            if i < len(lines) and lines[i].strip().startswith('|'):
                continue
            else:
                # End of table
                add_table_from_md(doc, table_lines)
                table_lines = []
                in_table = False
            continue
        elif in_table:
            add_table_from_md(doc, table_lines)
            table_lines = []
            in_table = False

        # ── Headings ─────────────────────────────────────────────────────────
        h4_match = re.match(r'^#{4}\s+(.*)', line)
        h3_match = re.match(r'^#{3}\s+(.*)', line)
        h2_match = re.match(r'^#{2}\s+(.*)', line)
        h1_match = re.match(r'^#\s+(.*)', line)

        if h1_match:
            if skip_first_h1:
                skip_first_h1 = False
            else:
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', h1_match.group(1))
                doc.add_heading(text, level=1)
            i += 1
            continue

        if h2_match:
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', h2_match.group(1))
            doc.add_heading(text, level=2)
            i += 1
            continue

        if h3_match:
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', h3_match.group(1))
            doc.add_heading(text, level=3)
            i += 1
            continue

        if h4_match:
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', h4_match.group(1))
            doc.add_heading(text, level=4)
            i += 1
            continue

        # ── Separator ────────────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '003878')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────────
        if line.strip().startswith('>'):
            text = line.strip().lstrip('> ').strip()
            if text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                run = p.add_run(text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            i += 1
            continue

        # ── Lists ────────────────────────────────────────────────────────────
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', line)
        if list_match:
            indent_str = list_match.group(1)
            bullet = list_match.group(2)
            text = list_match.group(3)
            level = len(indent_str) // 2
            is_numbered = bool(re.match(r'\d+\.', bullet))

            style = 'List Number' if is_numbered else 'List Bullet'
            try:
                p = doc.add_paragraph(style=style)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)

            apply_inline_formatting(p, text)
            i += 1
            continue

        # ── Empty line ───────────────────────────────────────────────────────
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # ── Normal paragraph ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        apply_inline_formatting(p, line)
        i += 1

    # Flush remaining table
    if table_lines:
        add_table_from_md(doc, table_lines)


# ─────────────────────────────────────────────────────────────────────────────
# GÉNÉRATION PRINCIPALE
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(deliverable: dict) -> bool:
    """Générer un fichier .docx à partir d'un fichier .md."""
    source_path = ROOT_DIR / deliverable["source"]
    output_path = OUTPUT_DIR / deliverable["output"]

    if not source_path.exists():
        print(f"  ⚠  Source introuvable : {source_path}")
        return False

    print(f"  ↪  {deliverable['source']} → {deliverable['output']}")

    # Lire le contenu Markdown
    with open(source_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Créer le document
    doc = Document()

    # Configurer les marges
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # Configurer les styles
    setup_document_styles(doc)

    # Page de titre
    add_title_page(doc, deliverable)

    # Corps du document
    md_to_docx(doc, md_content)

    # Pied de page
    for section in doc.sections:
        footer = section.footer
        p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p_footer.clear()
        p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_footer.add_run(
            f"{AUTHOR} — {PROJECT} — {DATE} — Document confidentiel"
        )
        run_f.font.name = 'Calibri'
        run_f.font.size = Pt(8)
        run_f.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Sauvegarder
    doc.save(str(output_path))
    return True


def main():
    """Point d'entrée principal."""
    print("=" * 70)
    print("  StatEduc Mobile — Génération des livrables Word")
    print(f"  Auteur : {AUTHOR}")
    print(f"  Date   : {DATE}")
    print("=" * 70)
    print()

    # Créer le répertoire de sortie
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Répertoire de sortie : {OUTPUT_DIR}")
    print()

    # Générer chaque livrable
    success = 0
    errors  = 0

    for deliverable in DELIVERABLES:
        ok = generate_docx(deliverable)
        if ok:
            success += 1
        else:
            errors += 1

    # Résumé
    print()
    print("=" * 70)
    print(f"  Résultat : {success} document(s) généré(s), {errors} erreur(s)")
    print(f"  Répertoire : {OUTPUT_DIR.resolve()}")
    print("=" * 70)

    if errors > 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
